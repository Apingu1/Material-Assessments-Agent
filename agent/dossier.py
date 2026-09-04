from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .evidence import EvidenceCapture


APPENDIX_FONT = "Arial"


def _font_run(run, size: float | None = None) -> None:
    run.font.name = APPENDIX_FONT
    if size is not None:
        run.font.size = Pt(size)


def _clean_title(title: str) -> str:
    return " ".join((title or "").split()).strip()


def _display_source(capture: EvidenceCapture) -> str:
    source = capture.source
    raw = f"{source.title} {source.publisher} {source.url}".lower()
    title = _clean_title(source.title)

    if "british pharmacopoeia" in raw:
        return "British Pharmacopoeia"
    if "european pharmacopoeia" in raw or "edqm" in raw:
        return "European Pharmacopoeia"
    if "bnf.nice.org.uk" in raw or "british national formulary" in raw:
        return "BNF/NICE"
    if "pubchem" in raw:
        return f"PubChem - {title}" if title and title.lower() != "pubchem" else "PubChem"
    if "pubmed" in raw or "pubmed.ncbi.nlm.nih.gov" in raw:
        return f"PubMed - {title}" if title else "PubMed"
    if "medicines.org.uk" in raw or "electronic medicines compendium" in raw:
        return f"eMC - {title}" if title else "eMC"
    if "dailymed" in raw:
        return f"DailyMed - {title}" if title else "DailyMed"
    if "hpra.ie" in raw or "irish medicines board" in raw:
        return f"HPRA - {title}" if title else "HPRA"
    if "cayman" in raw and "safety data sheet" in raw:
        material = re.sub(r"^safety data sheet:\s*", "", title, flags=re.IGNORECASE)
        return f"Cayman Chemical - {material} Safety Data Sheet".strip()
    return title or source.publisher or "Source"


def _add_scaled_image(doc: Document, path: Path) -> None:
    max_width = 6.05
    max_height = 6.55
    with Image.open(path) as image:
        width_px, height_px = image.size
    if not width_px or not height_px:
        return
    aspect = width_px / height_px
    if aspect >= max_width / max_height:
        width_in = max_width
        height_in = width_in / aspect
    else:
        height_in = max_height
        width_in = height_in * aspect

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(0)
    para.add_run().add_picture(str(path), width=Inches(width_in), height=Inches(height_in))


def append_evidence_dossier(
    filled_form: Path,
    output_path: Path,
    captures: list[EvidenceCapture],
) -> None:
    doc = Document(filled_form)

    for index, capture in enumerate(captures):
        # The controlled template already flows onto the next page after page 3.
        # Avoid an extra initial page break, which previously created a blank page 4.
        if index > 0:
            doc.add_page_break()

        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_after = Pt(5)
        run = heading.add_run(f"APPENDIX {capture.appendix_label} - {capture.group} Evidence")
        run.bold = True
        _font_run(run, 13)

        _labelled_paragraph(doc, "Source", _display_source(capture))
        _labelled_paragraph(doc, "URL", capture.source.url, size=8.5)
        _labelled_paragraph(doc, "Relevant finding", capture.source.relevant_extract)
        _labelled_paragraph(doc, "Interpretation", capture.source.interpretation)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("Evidence:")
        r.bold = True
        _font_run(r, 9.5)

        if capture.capture_path and capture.capture_path.exists():
            _add_scaled_image(doc, capture.capture_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def _labelled_paragraph(doc: Document, label: str, value: str, size: float = 9.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    _font_run(label_run, size)
    value_run = p.add_run(value or "N/A")
    _font_run(value_run, size)


def find_soffice() -> str | None:
    candidates = [
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path | None:
    soffice = find_soffice()
    if not soffice:
        return None
    output_dir.mkdir(parents=True, exist_ok=True)
    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if result.returncode != 0:
        return None
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    return pdf_path if pdf_path.exists() else None
