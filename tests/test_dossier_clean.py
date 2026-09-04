from pathlib import Path

from docx import Document
from PIL import Image

from agent.dossier import _display_source, append_evidence_dossier
from agent.evidence import EvidenceCapture, _page_matches_source
from agent.models import EvidenceApplicability, EvidenceSource, SourceType


def _source():
    return EvidenceSource(
        title="Haloperidol - British Pharmacopoeia",
        publisher="British Pharmacopoeia (reproduced by Trung Tâm Thuốc)",
        url="https://example.com/bp.pdf",
        tier=1,
        source_type=SourceType.PHARMACOPOEIAL,
        relevant_extract="Practically insoluble in water, slightly soluble in ethanol.",
        interpretation="Haloperidol is practically insoluble in water.",
        applicability=[EvidenceApplicability.CHEMICAL_SPECIES],
    )


def test_appendix_source_name_is_human_clean():
    capture = EvidenceCapture(_source(), "Water Solubility", "3A", None, None, "CAPTURED", "note")
    assert _display_source(capture) == "British Pharmacopoeia"


def test_page_match_rejects_unrelated_forbidden_page_and_accepts_evidence():
    source = _source()
    assert not _page_matches_source("403 Forbidden", source)
    assert _page_matches_source("Haloperidol is practically insoluble in water and slightly soluble in ethanol.", source)


def test_visible_appendix_omits_ai_metadata_and_keeps_url(tmp_path: Path):
    form = tmp_path / "form.docx"
    Document().save(form)
    image_path = tmp_path / "evidence.png"
    Image.new("RGB", (900, 600), "white").save(image_path)
    capture = EvidenceCapture(
        _source(),
        "Water Solubility",
        "3A",
        image_path,
        None,
        "CAPTURED",
        "Relevant PDF page captured",
    )
    out = tmp_path / "dossier.docx"
    append_evidence_dossier(form, out, [capture])
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Source: British Pharmacopoeia" in text
    assert "URL: https://example.com/bp.pdf" in text
    assert "Relevant finding:" in text
    assert "Interpretation:" in text
    assert "Evidence:" in text
    assert "Publisher:" not in text
    assert "Source tier:" not in text
    assert "Evidence capture:" not in text
    assert "Capture note:" not in text
    assert "reproduced by" not in text
